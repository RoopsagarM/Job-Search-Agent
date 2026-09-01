"""
CrewAI-based resume tailoring — real multi-agent orchestration (a Writer
agent and an ATS Reviewer agent, coordinated through CrewAI's Crew), with
the write -> review -> revise-if-needed loop controlled by plain Python
instead of CrewAI's Flow class.

Flow was dropped deliberately after causing two real, measured problems:
  1. It auto-creates a LanceDB-backed Memory object on every instance
     (a feature we don't use), adding several real seconds of overhead
     to EVERY resume generation, not just the first.
  2. Its own tracing/telemetry layer behaves independently of Crew's,
     and even with tracing explicitly disabled kept attempting network
     calls.
CrewAI's core building blocks (Agent, Task, Crew) don't have either
problem — this still satisfies "use CrewAI" and still orchestrates two
agents, it's just faster and far more predictable without Flow in the
middle.
"""

import json
import re
import time

from crewai import LLM, Agent, Task, Crew, Process

DEFAULT_MODEL = "anthropic/claude-sonnet-5"
DEFAULT_REVIEWER_MODEL = "anthropic/claude-haiku-4-5-20251001"

WRITER_BACKSTORY = (
    "You are an expert technical resume writer who specializes in "
    "tailoring resumes to specific job descriptions while keeping "
    "every claim truthful. You reorder and reword content to emphasize "
    "genuinely relevant experience, mirror the job's key terms where "
    "truthfully applicable, and format for both human readers and "
    "ATS (Applicant Tracking System) parsers: clear section headers, "
    "no tables or graphics, standard fonts implied, keywords from the "
    "job description used naturally. You never invent employers, "
    "titles, skills, or accomplishments not supported by the original "
    "resume."
)

REVIEWER_BACKSTORY = (
    "You are an ATS (Applicant Tracking System) compliance and resume "
    "quality reviewer. You score resumes exactly as an ATS parser and "
    "recruiter screener would: keyword match against the job "
    "description, clean section structure, quantified achievements, "
    "no formatting that breaks ATS parsing, appropriate length. You "
    "are specific and actionable in your feedback."
)


def _strip_json_fences(text: str) -> str:
    text = text.strip()
    if text.startswith("```"):
        text = text.split("```")[1]
        if text.startswith("json"):
            text = text[4:]
    return text.strip()


def _parse_review_json(raw_text: str) -> dict:
    """
    Tolerates the common ways a model deviates from "return ONLY JSON"
    (stray commentary, partial markdown fences). Never raises — a
    Reviewer hiccup should never destroy an already-completed, already
    paid-for Writer draft. Falls back to a neutral score if nothing
    parseable is found.
    """
    candidate = _strip_json_fences(raw_text)
    try:
        return json.loads(candidate)
    except json.JSONDecodeError:
        pass

    match = re.search(r"\{.*\}", raw_text, flags=re.DOTALL)
    if match:
        try:
            return json.loads(match.group(0))
        except json.JSONDecodeError:
            pass

    return {
        "ats_score": 50,
        "feedback": "Reviewer response couldn't be parsed this round — "
                    "showing the draft anyway. Try Reject & regenerate for a fresh score.",
    }


def _build_crew(anthropic_api_key: str, model: str, reviewer_model: str = None):
    """
    reviewer_model defaults to a faster model (Haiku) than the writer —
    scoring/feedback is a much simpler judgment task than writing the
    resume itself, so this cuts real wall-clock time per round.
    """
    reviewer_model = reviewer_model or DEFAULT_REVIEWER_MODEL

    # Note: the current Anthropic SDK's Messages.create() no longer accepts
    # `temperature` (or top_p/top_k) — removed in a breaking SDK update.
    # timeout is set explicitly because CrewAI's LLM defaults to no timeout
    # at all — a stalled network call would otherwise hang indefinitely
    # with the UI showing nothing but a spinner and no error.
    writer_llm = LLM(model=model, api_key=anthropic_api_key, timeout=60)
    reviewer_llm = LLM(model=reviewer_model, api_key=anthropic_api_key, timeout=60)

    writer = Agent(
        role="ATS Resume Writer",
        goal="Produce a truthful, tailored, ATS-optimized resume for the target job.",
        backstory=WRITER_BACKSTORY,
        llm=writer_llm,
        verbose=False,
    )
    reviewer = Agent(
        role="ATS Reviewer",
        goal="Score the resume's ATS-friendliness and fit, and give concrete improvement feedback.",
        backstory=REVIEWER_BACKSTORY,
        llm=reviewer_llm,
        verbose=False,
    )

    write_task = Task(
        description=(
            "ORIGINAL RESUME:\n{resume}\n\n"
            "TARGET JOB TITLE: {job_title}\n"
            "TARGET COMPANY: {company}\n"
            "JOB DESCRIPTION:\n{description}\n\n"
            "{revision_note}\n\n"
            "Rewrite the resume tailored to this job. Return ONLY the "
            "resume text — no commentary, no markdown fences."
        ),
        expected_output="The complete tailored resume text.",
        agent=writer,
    )
    review_task = Task(
        description=(
            "Review the resume produced in the previous task against this "
            "job description:\n{description}\n\n"
            "Respond with ONLY a JSON object (no markdown fences) in this "
            'exact form: {{"ats_score": <0-100 int>, "feedback": '
            '"<2-4 sentences of specific, actionable improvement feedback>"}}'
        ),
        expected_output="A JSON object with ats_score and feedback.",
        agent=reviewer,
        context=[write_task],
    )

    crew = Crew(
        agents=[writer, reviewer],
        tasks=[write_task, review_task],
        process=Process.sequential,
        verbose=False,
        tracing=False,
    )
    return crew, write_task


def _run_one_round(cv_text, job, anthropic_api_key, model, reviewer_model,
                    previous_resume=None, previous_feedback=None):
    """
    One Write -> Review pass. Returns (resume_text, ats_score, feedback).
    Retries once on a transient connection/timeout error before giving up —
    a brief network blip shouldn't force a manual re-click.
    """
    revision_note = ""
    if previous_resume and previous_feedback:
        revision_note = (
            f"This is a REVISION. Your previous draft was:\n{previous_resume}\n\n"
            f"A reviewer gave this feedback — address it directly:\n{previous_feedback}"
        )

    inputs = {
        "resume": cv_text,
        "job_title": job.get("title", ""),
        "company": job.get("company", ""),
        "description": job.get("description", ""),
        "revision_note": revision_note,
    }

    last_error = None
    for attempt in range(2):
        try:
            crew, write_task = _build_crew(anthropic_api_key, model, reviewer_model)
            result = crew.kickoff(inputs=inputs)
            resume_text = write_task.output.raw.strip()
            review = _parse_review_json(result.raw)
            return resume_text, review.get("ats_score", 0), review.get("feedback", "")
        except Exception as e:
            error_name = type(e).__name__
            is_connection_issue = "Connection" in error_name or "Timeout" in error_name
            last_error = e
            if attempt == 0 and is_connection_issue:
                time.sleep(3)
                continue
            raise
    raise last_error


def generate_with_orchestration(cv_text: str, job: dict, anthropic_api_key: str,
                                 model: str = DEFAULT_MODEL,
                                 reviewer_model: str = None,
                                 min_score: int = 65, max_attempts: int = 1) -> dict:
    """
    Runs the write -> review -> (revise if needed) loop. The loop itself
    is a plain Python for-loop (not CrewAI Flow) — it stops early once
    the score clears min_score, or after max_attempts rounds, whichever
    comes first. Two real CrewAI agents (Writer, Reviewer) do the actual
    work each round; this function just decides whether to run another
    round, which is the orchestration logic.
    Returns {"resume_text", "ats_score", "feedback", "attempts_used", "history"}.
    """
    previous_resume, previous_feedback = None, None
    history = []
    resume_text, ats_score, feedback = "", 0, ""

    for attempt in range(1, max_attempts + 1):
        resume_text, ats_score, feedback = _run_one_round(
            cv_text, job, anthropic_api_key, model, reviewer_model,
            previous_resume, previous_feedback,
        )
        history.append({"attempt": attempt, "ats_score": ats_score})
        if ats_score >= min_score:
            break
        previous_resume, previous_feedback = resume_text, feedback

    return {
        "resume_text": resume_text,
        "ats_score": ats_score,
        "feedback": feedback,
        "attempts_used": len(history),
        "history": history,
    }


def generate_or_revise(cv_text: str, job: dict, anthropic_api_key: str,
                        model: str = DEFAULT_MODEL, reviewer_model: str = None,
                        previous_resume: str = None,
                        previous_feedback: str = None) -> dict:
    """Single manual pass — used by the UI's 'Reject & regenerate' button."""
    resume_text, ats_score, feedback = _run_one_round(
        cv_text, job, anthropic_api_key, model, reviewer_model,
        previous_resume, previous_feedback,
    )
    return {"resume_text": resume_text, "ats_score": ats_score, "feedback": feedback}


def build_docx_bytes(resume_text: str, heading: str) -> bytes:
    """Turn tailored resume text into a downloadable/saveable .docx file's bytes."""
    import io
    import docx as docx_lib

    doc = docx_lib.Document()
    doc.add_heading(heading, level=1)
    for line in resume_text.split("\n"):
        if line.strip():
            doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
